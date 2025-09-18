import os
import re
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union

# Importuri pentru diferite tipuri de fișiere
import pdfplumber
from docxtpl import DocxTemplate
from docx import Document
import pytesseract
import cv2
import numpy as np
from pdf2image import convert_from_path


# Configurare logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Clasă pentru procesarea diferitelor tipuri de documente."""

    def __init__(self):
        # Configurare Tesseract pentru OCR
        self._configure_tesseract()
        # Tipuri de fișiere suportate
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}

        # Regex-uri îmbunătățite pentru extragerea datelor
        self.patterns = self._init_patterns()


    def _configure_tesseract(self):
        """Configurează calea către Tesseract OCR."""
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            rf"C:\Users\{os.getenv('USERNAME', '')}\AppData\Local\Tesseract-OCR\tesseract.exe",
            "tesseract"  # dacă e în PATH
        ]
        for path in possible_paths:
            if os.path.exists(path) or path == "tesseract":
                pytesseract.pytesseract.tesseract_cmd = path
                logger.info(f"Configurat Tesseract: {path}")
                return
        logger.warning("Nu s-a găsit Tesseract OCR, verifică instalarea.")

    def _init_patterns(self) -> Dict[str, List[str]]:
        """Inițializează pattern-urile regex pentru diferite tipuri de date."""
        return {
            'nume_prenume': [
                r"(?:Asociat(?:ul)?\s+unic\s*[:\-]?\s*)([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\s\-\']{3,60}?)(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"(?:Administrator(?:ul)?\s*[:\-]?\s*)([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\s\-\']{3,60}?)(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"(?:Beneficiarul\s+real\s+.*?este\s*[:\-]?\s*)([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\s\-\']{3,60}?)(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"(?:Administrarea\s+societății\s+se\s+face\s+de\s+către\s*[:\-]?\s*)([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\s\-\']{3,60}?)(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"(?:Controlul\s+.*?exercită.*?calitate\s+de\s+asociat\s+unic,?\s*)([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\s\-\']{3,60}?)(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"([A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț]{2,20}(?:\s+[A-ZĂÂÎÎȘȚ][A-ZĂÂÎÎȘȚa-zăâîîșț\-]{2,20}){1,3})(?=,\s*(?:cetatean|cetăţean|născut|nascut))",
                r"(?:n[ăa]scut[ăa]?\s+(?:la\s+data\s+de\s+|la\s+))(\d{1,2}\s+(?:ianuarie|februarie|martie|aprilie|mai|iunie|iulie|august|septembrie|octombrie|noiembrie|decembrie)\s+\d{4})",
                r"(?:Birth\s+date[:\s]*)(\d{1,2}\s+[A-Za-z]+\s+\d{4})"
            ],
            'data_nasterii': [
                r"(?:n[ăa]scut[ăa]?\s+la\s+data\s+de\s+)(\d{1,2}\.\d{1,2}\.\d{4})",
                r"(?:născut[ăa]?\s+la\s+)(\d{1,2}\.\d{1,2}\.\d{4})",
                r"(?:Data\s+nașterii[:\s]*)(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})",
                r"(\d{1,2}\.\d{1,2}\.\d{4})(?=,\s*(?:in|în)\s+(?:Mun|Com|Sat|Jud))",
                r"(?:Born\s+on[:\s]*)(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})"
            ],
            'domiciliu': [
                r"domiciliat[ăa]?\s+(?:în|in)\s+(.+?)(?=,\s*(?:identificat|având|posesor|CNP|cu\s+domiciliul))",
                r"domiciliată\s+(?:în|in)\s+(.+?)(?=,\s*(?:identificat|având|posesor|CNP))",
                r"(?:cu\s+domiciliul\s+(?:în|in|la)\s+)(.+?)(?=,\s*(?:identificat|având|posesor|CNP))",
                r"(?:Domiciliu[:\s]+)(.+?)(?=,\s*(?:identificat|având|posesor|CNP|telefon))",
                r"(?:Adres[ăa][:\s]+)(.+?)(?=,\s*(?:tel|CNP|identificat))",
                r"(?:Residence[:\s]+)(.+?)(?=,\s*(?:identified|CNP))"
            ],
            'tip_act_identificare': [
                r"(CI\s+seria\s+[A-Z]{1,3},?\s*nr\.?\s*\d{6,10})",
                r"(BI\s+seria\s+[A-Z]{1,3},?\s*nr\.?\s*\d{6,10})",
                r"(Carte\s+de\s+identitate\s+seria\s+[A-Z]{1,3},?\s*nr\.?\s*\d{6,10})",
                r"identificat[ăa]?\s+cu\s+(CI\s+seria\s+[A-Z]{1,3},?\s*nr\.?\s*\d{6,10})",
                r"(Pașaport\s+nr\.?\s*[\w\d]{6,15})",
                r"(ID\s+(?:card|nr\.?)\s*[\w\d]{3,15})",
                r"(Passport\s+(?:nr\.?|no\.?)\s*[\w\d]{3,15})"
            ],
            'CNP': [
                r"CNP\s*[:\-]?\s*([0-9]{13})",
                r"(?:Cod\s+numeric\s+personal[:\s]*)([0-9]{13})",
                r"(?:Personal\s+code[:\s]*)([0-9]{13})"
            ],
            'nume_societate': [
                r"(?:Denumirea\s+societ[ăa]ții\s+este\s*[:\-]?\s*)([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+S\.R\.L\.)",
                r"(?:Societății\s+)([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+S\.R\.L\.)",
                r"Societatea\s+([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+S\.R\.L\.)",
                r"înființarea\s+Societății\s+([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+S\.R\.L\.)",
                r"([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+S\.R\.L\.)(?=\s*,?\s*conform)",
                r"(?:Firm[ăa][:\s]+)([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+(?:S\.R\.L\.|SRL|S\.A\.|SA|S\.C\.|SC))",
                r"(?:Company\s+name[:\s]+)([A-Z0-9\s\.\-,ĂÂÎÎȘȚăâîîșț]+(?:S\.R\.L\.|SRL|S\.A\.|SA|S\.C\.|SC))"
            ],
            'sediu_firma': [
                r"(?:Sediul\s+societ[ăa]ții\s+este\s+(?:în|in|la)\s+)(.+?)(?=\.|Art\.|CAPITOLUL|\n)",
                r"(?:Sediu\s+social\s*[:\-]?\s*)(.*?)(?=\n|,\s*(?:Telefon|Tel|Email|CUI|Cod))",
                r"(?:Sediul[:\s]+)(.*?)(?=\n|,\s*(?:Telefon|Tel|Email|CUI|Cod))",
                r"(?:Headquarters[:\s]+)(.*?)(?=\n|,\s*(?:Phone|Tel|Email|CUI|Code))",
                r"(?:cu\s+sediul\s+social\s+în\s+)(.+?)(?=,|\n|$)"
            ],
            'CUI': [
                r"(?:Cod\s+Unic\s+de\s+Înregistrare\s*[:\-]?\s*(?:RO)?\s*)(\d{2,10})",
                r"(?:CUI\s*[:\-]?\s*(?:RO)?\s*)(\d{2,10})",
                r"(?:C\.?U\.?I\.?[:\s]*)(\d{2,10})",
                r"(?:VAT\s+number[:\s]*)(\d{2,10})"
            ],
            'data_inregistrarii': [
                r"(?:din\s+data\s+de\s+)(\d{1,2}[./-]\d{1,2}[./-]\d{4})",
                r"(?:Din.)(\d{1,2}[./-]\d{1,2}[./-]\d{4})",
                r"(?:înregistrat[ăa]?\s+la\s+data\s+de\s+)(\d{1,2}[./-]\d{1,2}[./-]\d{4})",
                r"(?:Registration\s+date[:\s]+)(\d{1,2}[./-]\d{1,2}[./-]\d{4})"
            ],
            'id_unic_european': [
                r"\(EUID\)\s*[:\-]?\s*([A-Z0-9.\-]{10,30})",
                r"(?:EUID[:\s]+)([A-Z0-9.\-]{10,30})"
            ],
            'numar_ordine': [
                r"(?:Nr\.?\s*(?:de)?\s*ordine\s*(?:în\s*registrul\s*comerțului)?\s*[:\-]?\s*)([A-Z0-9/]{1,20})",
                r"(?:Trade\s+register\s+number[:\s]+)([A-Z0-9/]{1,20})"
            ],
            'sediu_societate': [
                r"(?:Sediul\s+societ[ăa]ții\s+este\s+(?:în|la)\s+)(.+?)(?=\.|Art\.|CAPITOLUL)",
                r"(?:Sediul\s+societății\s+este\s+în\s+)(.+?)(?=\n|\.)"
            ]
        }
    def safe_search(self, patterns: List[str], text: str) -> str:
        """
        Caută prin mai multe pattern-uri regex și returnează prima potrivire găsită.
        """
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                # Încearcă să returneze primul grup de captură, altfel întoarce întreaga potrivire
                try:
                    result = match.group(1).strip() if match.groups() else match.group(0).strip()
                    # Pentru CUI, concatenează RO cu cifra dacă există ambele grupuri
                    if len(match.groups()) > 1 and match.group(1) and match.group(2):
                        result = f"{match.group(1)}{match.group(2)}"
                    if result:
                        return self._clean_extracted_text(result)
                except (IndexError, AttributeError):
                    continue
        return ""

    def _clean_extracted_text(self, text: str) -> str:
        """Curăță textul extras de caractere nedorite."""
        # Elimină caractere speciale și spații multiple
        text = re.sub(r'[^\w\s\.\-/,ĂÂÎÎȘȚăâîîșț]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def preprocess_image(self, image_path: str) -> np.ndarray:
        """Preprocesează imaginea pentru OCR mai bun."""
        image = cv2.imread(image_path)

        # Conversie la grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Aplicare filtru pentru reducerea noise-ului
        denoised = cv2.medianBlur(gray, 5)

        # Îmbunătățire contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)

        # Binarizare adaptivă
        binary = cv2.adaptiveThreshold(enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                       cv2.THRESH_BINARY, 11, 2)

        return binary

    def read_pdf(self, path: str) -> str:
        """Citește textul din PDF (nativ sau scanat prin OCR)."""
        text = ""
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Eroare la citirea PDF {path}: {e}")

        # Dacă nu s-a extras text -> fallback OCR
        if not text.strip():
            logger.info(f"PDF {path} pare scanat, aplic OCR...")
            try:
                images = convert_from_path(path)
                for i, img in enumerate(images, start=1):
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        img.save(tmp.name, "PNG")
                        try:
                            # Preprocesează imaginea înainte de OCR
                            processed = self.preprocess_image(tmp.name)
                            config = r"--oem 3 --psm 6 -l ron+eng"
                            page_text = pytesseract.image_to_string(processed, config=config)
                            if page_text.strip():
                                text += page_text + "\n"
                        except Exception as ocr_err:
                            logger.error(f"Eroare OCR la pagina {i} din {path}: {ocr_err}")
                        finally:
                            os.unlink(tmp.name)
            except Exception as e:
                logger.error(f"Eroare la conversia PDF {path} în imagini pentru OCR: {e}")

        return text

    def read_docx(self, path: str) -> str:
        """Citește textul dintr-un DOCX."""
        try:
            doc = Document(path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Citește și din tabele
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += f" {cell.text}"

            return text
        except Exception as e:
            logger.error(f"Eroare la citirea DOCX {path}: {e}")
            return ""

    def read_image_ocr(self, path: str) -> str:
        """Citește textul dintr-o imagine folosind OCR, cu preprocesare și fallback."""
        try:
            # Preprocesare imagine pentru OCR mai bun
            processed = self.preprocess_image(path)

            # Config inițial (română + engleză)
            config = r"--oem 3 --psm 6 -l ron+eng"
            text = pytesseract.image_to_string(processed, config=config).strip()

            # Dacă nu a returnat nimic, încearcă fallback doar pe engleză
            if not text:
                logger.debug(f"OCR cu 'ron+eng' nu a returnat rezultate pentru {path}, încerc fallback ENG.")
                fallback_config = r"--oem 3 --psm 6 -l eng"
                text = pytesseract.image_to_string(processed, config=fallback_config).strip()

            # Normalizează textul
            text = re.sub(r"\s+", " ", text)  # elimină spații multiple
            text = text.strip()

            return text

        except Exception as e:
            logger.error(f"Eroare la OCR pentru {path}: {e}")
            return ""

    def read_file(self, path: str) -> str:
        """Citește conținutul unui fișier în funcție de extensie."""
        path_obj = Path(path)
        extension = path_obj.suffix.lower()

        if extension == '.pdf':
            return self.read_pdf(path)
        elif extension in ['.docx', '.doc']:
            return self.read_docx(path)
        elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self.read_image_ocr(path)
        else:
            logger.warning(f"Extensie nesuportată: {extension}")
            return ""

    def extract_data(self, text: str) -> Dict[str, str]:
        """
        Extrage datele dintr-un text, indiferent de tipul documentului.
        Rulează toate pattern-urile și returnează doar câmpurile găsite.
        """
        data = {}
        for field, patterns in self.patterns.items():
            value = self.safe_search(patterns, text)
            if value:
                data[field] = value
        return data

    def process_directory(self, input_dir: str, template_path: str, output_path: str):
        all_data: Dict[str, List[str]] = {}

        for file_path in Path(input_dir).iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                logger.info(f"Procesez fișierul: {file_path}")
                text = self.read_file(str(file_path))
                extracted = self.extract_data(text)

                for k, v in extracted.items():
                    all_data.setdefault(k, []).append(v)

        logger.info(f"Date totale extrase: { {k: len(v) for k, v in all_data.items()} }")

        if all_data:
            # folosim prima valoare găsită pentru fiecare câmp
            context = {k: v[0] for k, v in all_data.items()}

            # completează câmpurile obligatorii lipsă
            required_fields = list(self.patterns.keys())
            context = self.prompt_missing_fields(context, required_fields)

            self.generate_procura(template_path, output_path, context)
        else:
            logger.warning("Nu s-au găsit date în niciun fișier.")

    def prompt_missing_fields(self, data: dict, required_fields: list) -> dict:
        """
        Verifică câmpurile obligatorii. Dacă lipsesc, cere utilizatorului să le completeze.
        - data: dict cu datele extrase
        - required_fields: lista câmpurilor obligatorii (ex: ['nume_prenume', 'CNP', 'tip_act_identificare', 'data_nasterii', 'domiciliu'])
        """
        for field in required_fields:
            if field not in data or not data[field]:
                user_input = input(f"Lipseste '{field}'. Te rog completează: ").strip()
                data[field] = user_input
        return data

    def generate_procura(self, template_path: str, output_path: str, context: Dict[str, str]):
        """Generează procura folosind template-ul și datele extrase."""
        try:
            doc = DocxTemplate(template_path)
            doc.render(context)
            doc.save(output_path)
            logger.info(f"Procura generată cu succes la: {output_path}")
        except Exception as e:
            logger.error(f"Eroare la generarea procurii: {e}")



def main():
    """Funcția principală."""
    # Configurare paths
    input_dir = r"C:\intrari\2"
    template_path = r"C:\model_procura\IMPUTERNICIRE_model_ro_eng.docx"
    output_path = r"C:\procuri_completate\PROCURA_2.docx"

    # Creează directoarele dacă nu există
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    processor = DocumentProcessor()
    processor.process_directory(input_dir, template_path, output_path)


if __name__ == "__main__":
    main()